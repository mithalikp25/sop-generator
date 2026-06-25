import re
import os
import streamlit as st
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from chain import get_initial_messages, get_next_question, generate_sop

load_dotenv()

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(page_title="SOP Generator", page_icon="📋", layout="centered")
st.title("📋 SOP Generator for Business Workflows")
st.caption("Describe your business process, answer 5 questions, get a complete SOP.")

# ── Session state ─────────────────────────────────────────────────────────────
for key, default in {
    "stage": "input",
    "messages": [],
    "question_num": 0,
    "current_question": "",
    "sop_output": "",
    "process_description": ""
}.items():
    if key not in st.session_state:
        st.session_state[key] = default


# ── Helpers ───────────────────────────────────────────────────────────────────
def render_inline(paragraph, text):
    """Write a line into a Word paragraph, converting **bold** markers."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part:
            paragraph.add_run(part)


def export_to_docx(sop_text):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    title_para = doc.add_heading('Standard Operating Procedure', 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    lines = sop_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Markdown table
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                if not re.match(r'^\|[-| :]+\|$', row):
                    table_lines.append(row)
                i += 1
            if table_lines:
                num_cols = len(table_lines[0].split("|")) - 2
                if num_cols > 0:
                    table = doc.add_table(rows=len(table_lines), cols=num_cols)
                    table.style = 'Table Grid'
                    for r_idx, row_text in enumerate(table_lines):
                        cells = [c.strip() for c in row_text.split("|")[1:-1]]
                        for c_idx in range(min(len(cells), num_cols)):
                            cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cells[c_idx])
                            p = table.cell(r_idx, c_idx).paragraphs[0]
                            run = p.add_run(cell_text)
                            if r_idx == 0:
                                run.bold = True
                    doc.add_paragraph()
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif re.match(r'^\d+[\.\)]\s', line):
            p = doc.add_paragraph(style='List Number')
            render_inline(p, re.sub(r'^\d+[\.\)]\s', '', line))
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            render_inline(p, line[2:])
        elif re.match(r'^\*[^*].*[^*]\*$', line):
            p = doc.add_paragraph()
            p.add_run(line.strip("*")).italic = True
        else:
            p = doc.add_paragraph()
            render_inline(p, line)

        i += 1

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Stage 1: Input ────────────────────────────────────────────────────────────
if st.session_state.stage == "input":
    st.markdown("### Step 1 — Describe your business process")
    process = st.text_area(
        "Write 2-3 sentences:",
        placeholder="e.g. We need to onboard new employees. This includes IT setup, HR paperwork, and team introduction.",
        height=120
    )
    if st.button("Start Interview →", type="primary"):
        if not process.strip():
            st.warning("Please describe your business process first.")
        else:
            with st.spinner("Starting interview..."):
                msgs = get_initial_messages(process)
                question = get_next_question(msgs)
                msgs.append({"role": "assistant", "content": question})
                st.session_state.messages = msgs
                st.session_state.current_question = question
                st.session_state.question_num = 1
                st.session_state.process_description = process
                st.session_state.stage = "interview"
                st.rerun()

# ── Stage 2: Interview ────────────────────────────────────────────────────────
elif st.session_state.stage == "interview":
    st.markdown("### Step 2 — Answer 5 clarifying questions")
    st.progress(st.session_state.question_num / 5,
                text=f"Question {st.session_state.question_num} of 5")

    for msg in st.session_state.messages:
        if msg["role"] == "assistant":
            with st.chat_message("assistant"):
                st.write(msg["content"])
        elif msg["role"] == "user" and msg != st.session_state.messages[1]:
            with st.chat_message("user"):
                st.write(msg["content"])

    answer = st.text_input(
        f"Your answer to Question {st.session_state.question_num}:",
        key=f"answer_{st.session_state.question_num}"
    )

    if st.button("Submit Answer →", type="primary"):
        if not answer.strip():
            st.warning("Please enter an answer.")
        else:
            st.session_state.messages.append({"role": "user", "content": answer})
            if st.session_state.question_num < 5:
                with st.spinner("Getting next question..."):
                    next_q = get_next_question(st.session_state.messages)
                    st.session_state.messages.append({"role": "assistant", "content": next_q})
                    st.session_state.current_question = next_q
                    st.session_state.question_num += 1
                    st.rerun()
            else:
                with st.spinner("Generating your SOP... please wait"):
                    sop = generate_sop(st.session_state.messages)
                    st.session_state.sop_output = sop
                    st.session_state.stage = "result"
                    st.rerun()

# ── Stage 3: Result ───────────────────────────────────────────────────────────
elif st.session_state.stage == "result":
    st.success("✅ Your SOP has been generated!")
    st.markdown("### 📄 Generated SOP")
    st.markdown(st.session_state.sop_output)
    st.divider()

    try:
        docx_buf = export_to_docx(st.session_state.sop_output)
        st.download_button(
            label="⬇️ Download as Word Document",
            data=docx_buf,
            file_name="sop_output.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        st.error(f"Word export error: {e}")
        st.download_button(
            label="⬇️ Download as Text File",
            data=st.session_state.sop_output,
            file_name="sop_output.txt",
            mime="text/plain"
        )

    if st.button("🔄 Generate Another SOP"):
        for key in ["stage", "messages", "question_num",
                    "current_question", "sop_output", "process_description"]:
            del st.session_state[key]
        st.rerun()